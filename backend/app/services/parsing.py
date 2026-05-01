import logging
import zipfile
from dataclasses import dataclass, field
from pathlib import Path, PurePosixPath
from typing import Any
from urllib.parse import unquote, urldefrag
from xml.etree import ElementTree

import fitz
from bs4 import BeautifulSoup
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class ParsedDocument:
    raw_text: str
    title: str | None = None
    author: str | None = None
    date: str | None = None
    language: str | None = None
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentParser:
    supported_source_types = {
        "pdf",
        "docx",
        "html",
        "epub",
        "txt",
        "md",
        "text",
        "log",
        "csv",
        "tsv",
        "json",
        "xml",
        "rst",
    }

    def parse(self, file_path: str | Path, source_type: str) -> ParsedDocument:
        normalized_source_type = source_type.lower().strip()
        path = Path(file_path)
        if normalized_source_type not in self.supported_source_types:
            raise ValueError(f"Unsupported source_type: {source_type}")
        if normalized_source_type == "pdf":
            return self._parse_pdf(path)
        if normalized_source_type == "docx":
            return self._parse_docx(path)
        if normalized_source_type == "html":
            return self._parse_html(path)
        if normalized_source_type == "epub":
            return self._parse_epub(path)
        return self._parse_plain_text(path, normalized_source_type)

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        text_parts: list[str] = []
        metadata: dict[str, Any] = {"source_type": "pdf", "page_count": 0}
        with fitz.open(path) as document:
            metadata["page_count"] = document.page_count
            metadata.update({key: value for key, value in document.metadata.items() if value})
            for page_index in range(document.page_count):
                try:
                    text_parts.append(document.load_page(page_index).get_text())
                except Exception as exc:
                    logger.warning("Failed to parse PDF page %s in %s: %s", page_index, path, exc)
        return ParsedDocument(
            raw_text="\n\n".join(text_parts).strip(),
            title=metadata.get("title") or path.stem,
            author=metadata.get("author"),
            date=metadata.get("creationDate") or metadata.get("modDate"),
            metadata=metadata,
        )

    def _parse_docx(self, path: Path) -> ParsedDocument:
        document = DocxDocument(path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        properties = document.core_properties
        metadata = {
            "source_type": "docx",
            "paragraph_count": len(paragraphs),
            "subject": properties.subject,
            "keywords": properties.keywords,
            "category": properties.category,
        }
        created = properties.created.isoformat() if properties.created else None
        return ParsedDocument(
            raw_text="\n\n".join(paragraphs).strip(),
            title=properties.title or path.stem,
            author=properties.author,
            date=created,
            metadata={key: value for key, value in metadata.items() if value is not None},
        )

    def _parse_html(self, path: Path) -> ParsedDocument:
        html = self._read_text(path)
        soup = BeautifulSoup(html, "html.parser")
        for element in soup(["script", "style", "noscript"]):
            element.decompose()
        title = soup.title.string.strip() if soup.title and soup.title.string else path.stem
        author = self._meta_content(soup, "author")
        language = soup.html.get("lang") if soup.html else None
        text = soup.get_text(separator="\n")
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        return ParsedDocument(
            raw_text="\n".join(lines),
            title=title,
            author=author,
            language=language,
            metadata={"source_type": "html"},
        )

    def _parse_epub(self, path: Path) -> ParsedDocument:
        with zipfile.ZipFile(path) as archive:
            rootfile_path = self._epub_rootfile_path(archive)
            package = ElementTree.fromstring(archive.read(rootfile_path))
            metadata = self._epub_metadata(package)
            item_paths = self._epub_spine_item_paths(package, rootfile_path)
            text_parts = [
                self._html_to_text(archive.read(item_path).decode("utf-8", errors="replace"))
                for item_path in item_paths
            ]

        raw_text = "\n\n".join(part for part in text_parts if part).strip()
        return ParsedDocument(
            raw_text=raw_text,
            title=metadata.get("title") or path.stem,
            author=metadata.get("author"),
            date=metadata.get("date"),
            language=metadata.get("language"),
            metadata={
                "source_type": "epub",
                "item_count": len(item_paths),
                **{key: value for key, value in metadata.items() if value},
            },
        )

    def _parse_plain_text(self, path: Path, source_type: str) -> ParsedDocument:
        return ParsedDocument(
            raw_text=self._read_text(path),
            title=path.stem,
            metadata={"source_type": source_type},
        )

    def _read_text(self, path: Path) -> str:
        try:
            return path.read_text(encoding="utf-8")
        except UnicodeDecodeError as exc:
            logger.warning("UTF-8 decode failed for %s, retrying with replacement characters: %s", path, exc)
            return path.read_text(encoding="utf-8", errors="replace")

    def _html_to_text(self, html: str) -> str:
        soup = BeautifulSoup(html, "html.parser")
        for element in soup(["script", "style", "noscript"]):
            element.decompose()
        lines = [
            line.strip()
            for line in soup.get_text(separator="\n").splitlines()
            if line.strip()
        ]
        return "\n".join(lines)

    def _epub_rootfile_path(self, archive: zipfile.ZipFile) -> str:
        container = ElementTree.fromstring(archive.read("META-INF/container.xml"))
        rootfile = container.find(".//{*}rootfile")
        full_path = rootfile.get("full-path") if rootfile is not None else None
        if not full_path:
            raise ValueError("EPUB container does not declare a rootfile")
        return full_path

    def _epub_spine_item_paths(self, package: ElementTree.Element, rootfile_path: str) -> list[str]:
        manifest = {
            item.get("id"): item.get("href")
            for item in package.findall(".//{*}manifest/{*}item")
            if item.get("id") and item.get("href")
        }
        spine_ids = [
            itemref.get("idref")
            for itemref in package.findall(".//{*}spine/{*}itemref")
            if itemref.get("idref")
        ]
        hrefs = [manifest[idref] for idref in spine_ids if idref in manifest]
        if not hrefs:
            hrefs = [
                item.get("href")
                for item in package.findall(".//{*}manifest/{*}item")
                if item.get("href")
                and item.get("media-type") in {"application/xhtml+xml", "text/html"}
            ]
        package_dir = PurePosixPath(rootfile_path).parent
        return [self._epub_item_path(package_dir, href) for href in hrefs if href]

    def _epub_metadata(self, package: ElementTree.Element) -> dict[str, str | None]:
        return {
            "title": self._epub_metadata_text(package, "title"),
            "author": self._epub_metadata_text(package, "creator"),
            "date": self._epub_metadata_text(package, "date"),
            "language": self._epub_metadata_text(package, "language"),
        }

    @staticmethod
    def _epub_item_path(package_dir: PurePosixPath, href: str) -> str:
        path = unquote(urldefrag(href).url)
        return str(package_dir / path) if str(package_dir) != "." else path

    @staticmethod
    def _epub_metadata_text(package: ElementTree.Element, name: str) -> str | None:
        element = package.find(f".//{{*}}metadata/{{*}}{name}")
        return element.text.strip() if element is not None and element.text else None

    @staticmethod
    def _meta_content(soup: BeautifulSoup, name: str) -> str | None:
        tag = soup.find("meta", attrs={"name": name})
        if not tag:
            return None
        content = tag.get("content")
        return content.strip() if isinstance(content, str) else None
